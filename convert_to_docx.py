from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Read the markdown file
with open('hosting.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Create a new Document
doc = Document()

# Set up default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('Hosting Proposal for Knight and Carter', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Parse and add content
lines = content.split('\n')
for line in lines:
    line = line.strip()
    if not line:
        doc.add_paragraph()
    elif line.startswith('###'):
        doc.add_heading(line.replace('###', '').strip(), 2)
    elif line.startswith('##'):
        doc.add_heading(line.replace('##', '').strip(), 1)
    elif line.startswith('#'):
        pass
    elif line.startswith('-'):
        p = doc.add_paragraph(line.replace('-', '').strip(), style='List Bullet')
    elif line.endswith(':'):
        p = doc.add_paragraph(line)
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)
    else:
        doc.add_paragraph(line)

# Save the document
doc.save('Hosting Proposal.docx')
print('Document created: Hosting Proposal.docx')
